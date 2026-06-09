#!/usr/bin/env python3
"""
Markdown-to-DOCX Converter with Mermaid Diagram Support.

Converts a Markdown file (including Mermaid diagrams) into a professionally
formatted Word (.docx) document.

Mermaid diagrams are rendered via the mermaid.ink public API and embedded
as images. If the API is unavailable, a placeholder is inserted instead.

Usage:
    python md_to_docx.py [input.md] [output.docx]
    
    # Default: converts design_thinking_adaptika.md
    python md_to_docx.py
"""

import re
import os
import sys
import base64
import urllib.request
import urllib.error
import tempfile
import json
from io import BytesIO

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# =========================================================================
# STYLE CONSTANTS
# =========================================================================
FONT_MAIN = 'Calibri'
FONT_MONO = 'Consolas'
FONT_SIZE_BODY = 11
FONT_SIZE_H1 = 22
FONT_SIZE_H2 = 16
FONT_SIZE_H3 = 13
FONT_SIZE_H4 = 12
FONT_SIZE_SMALL = 9
FONT_SIZE_TABLE = 10

COLOR_PRIMARY = RGBColor(0x1A, 0x56, 0xDB)  # Blue
COLOR_HEADING = RGBColor(0x1E, 0x29, 0x3B)   # Dark navy
COLOR_TEXT = RGBColor(0x1F, 0x2A, 0x37)       # Near-black
COLOR_MUTED = RGBColor(0x6B, 0x72, 0x80)      # Gray
COLOR_ACCENT = RGBColor(0x05, 0x96, 0x69)     # Green
COLOR_TABLE_HEADER_BG = '1A56DB'               # For XML
COLOR_TABLE_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_TABLE_ALT_ROW = 'F3F4F6'
COLOR_BLOCKQUOTE_BAR = RGBColor(0xD1, 0xD5, 0xDB)
COLOR_CODE_BG = 'F3F4F6'

# Alert colors (GitHub-style)
ALERT_COLORS = {
    'NOTE':      ('E8F0FE', '1A56DB', 'ℹ️'),
    'TIP':       ('E6F4EA', '059669', '💡'),
    'IMPORTANT': ('FEF3C7', 'D97706', '⚠️'),
    'WARNING':   ('FFF7ED', 'EA580C', '⚠️'),
    'CAUTION':   ('FEE2E2', 'DC2626', '🔴'),
}

PAGE_MARGIN = Cm(2.54)  # 1 inch


# =========================================================================
# MERMAID RENDERING
# =========================================================================

def render_mermaid_via_mmdc(mermaid_code: str, output_path: str) -> bool:
    """Render Mermaid diagram using local mmdc (mermaid-cli via npx)."""
    import subprocess
    try:
        # Write mermaid code to temp file
        temp_mmd = output_path.replace('.png', '.mmd')
        with open(temp_mmd, 'w', encoding='utf-8') as f:
            f.write(mermaid_code)
        
        # Try running mmdc via npx
        result = subprocess.run(
            ['npx', '-y', '@mermaid-js/mermaid-cli', 
             '-i', temp_mmd, '-o', output_path, 
             '-b', 'white', '-t', 'default', '-s', '2'],
            capture_output=True, text=True, timeout=60
        )
        
        # Cleanup temp file
        if os.path.exists(temp_mmd):
            os.remove(temp_mmd)
        
        if result.returncode == 0 and os.path.exists(output_path):
            return True
        else:
            print(f"  ⚠️  mmdc error: {result.stderr[:200] if result.stderr else 'unknown'}")
            return False
    except Exception as e:
        print(f"  ⚠️  mmdc failed: {e}")
        return False


def render_mermaid_via_api(mermaid_code: str, output_path: str) -> bool:
    """Render Mermaid diagram via mermaid.ink public API (sanitized)."""
    try:
        # Sanitize: remove HTML tags that mermaid.ink doesn't support
        sanitized = re.sub(r'<br\s*/?>', '\\n', mermaid_code)
        sanitized = re.sub(r'<[^>]+>', '', sanitized)
        
        encoded = base64.urlsafe_b64encode(sanitized.encode('utf-8')).decode('ascii')
        url = f"https://mermaid.ink/img/{encoded}?bgColor=white&theme=default&scale=2"
        
        req = urllib.request.Request(url)
        req.add_header('User-Agent', 'ADAPTIKA-DOCX-Generator/1.0')
        
        with urllib.request.urlopen(req, timeout=30) as response:
            data = response.read()
            with open(output_path, 'wb') as f:
                f.write(data)
            return True
    except Exception as e:
        print(f"  ⚠️  API rendering failed: {e}")
        return False


def render_mermaid_to_png(mermaid_code: str, output_path: str) -> bool:
    """Render Mermaid diagram to PNG. Tries mmdc first, then API fallback."""
    # Method 1: Local mmdc (supports HTML tags in labels)
    print(f"    → Trying local mmdc...")
    if render_mermaid_via_mmdc(mermaid_code, output_path):
        return True
    
    # Method 2: mermaid.ink API (with sanitized code)
    print(f"    → Trying mermaid.ink API...")
    if render_mermaid_via_api(mermaid_code, output_path):
        return True
    
    return False


# =========================================================================
# MARKDOWN PARSER
# =========================================================================

class MarkdownParser:
    """Parse markdown into structured blocks for DOCX generation."""
    
    def __init__(self, md_text: str):
        self.lines = md_text.split('\n')
        self.blocks = []
        self.pos = 0
        self._parse()
    
    def _parse(self):
        while self.pos < len(self.lines):
            line = self.lines[self.pos]
            
            # Blank line
            if not line.strip():
                self.pos += 1
                continue
            
            # Horizontal rule
            if re.match(r'^---+\s*$', line.strip()):
                self.blocks.append({'type': 'hr'})
                self.pos += 1
                continue
            
            # Headings
            hm = re.match(r'^(#{1,6})\s+(.+)', line)
            if hm:
                level = len(hm.group(1))
                text = hm.group(2).strip()
                self.blocks.append({'type': 'heading', 'level': level, 'text': text})
                self.pos += 1
                continue
            
            # Mermaid code block
            if line.strip().startswith('```mermaid'):
                self.pos += 1
                mermaid_lines = []
                while self.pos < len(self.lines) and not self.lines[self.pos].strip().startswith('```'):
                    mermaid_lines.append(self.lines[self.pos])
                    self.pos += 1
                self.pos += 1  # skip closing ```
                self.blocks.append({'type': 'mermaid', 'code': '\n'.join(mermaid_lines)})
                continue
            
            # Generic code block  
            if line.strip().startswith('```'):
                lang = line.strip()[3:].strip()
                self.pos += 1
                code_lines = []
                while self.pos < len(self.lines) and not self.lines[self.pos].strip().startswith('```'):
                    code_lines.append(self.lines[self.pos])
                    self.pos += 1
                self.pos += 1  # skip closing ```
                self.blocks.append({'type': 'code', 'lang': lang, 'code': '\n'.join(code_lines)})
                continue
            
            # Alert blockquote (> [!NOTE] etc.)
            if re.match(r'^>\s*\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]', line):
                alert_match = re.match(r'^>\s*\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]', line)
                alert_type = alert_match.group(1)
                self.pos += 1
                alert_lines = []
                while self.pos < len(self.lines):
                    aline = self.lines[self.pos]
                    if aline.startswith('>'):
                        content = aline[1:].strip()
                        if content:
                            alert_lines.append(content)
                        self.pos += 1
                    else:
                        break
                self.blocks.append({
                    'type': 'alert', 
                    'alert_type': alert_type, 
                    'text': ' '.join(alert_lines)
                })
                continue
            
            # Regular blockquote
            if line.startswith('>'):
                bq_lines = []
                while self.pos < len(self.lines) and self.lines[self.pos].startswith('>'):
                    content = self.lines[self.pos][1:].strip()
                    bq_lines.append(content)
                    self.pos += 1
                self.blocks.append({'type': 'blockquote', 'text': '\n'.join(bq_lines)})
                continue
            
            # Table
            if '|' in line and self.pos + 1 < len(self.lines) and re.match(r'^[\s|:-]+$', self.lines[self.pos + 1]):
                table_lines = []
                while self.pos < len(self.lines) and '|' in self.lines[self.pos]:
                    table_lines.append(self.lines[self.pos])
                    self.pos += 1
                self._parse_table(table_lines)
                continue
            
            # Unordered list
            if re.match(r'^(\s*)[\*\-]\s+', line):
                list_items = []
                while self.pos < len(self.lines):
                    lm = re.match(r'^(\s*)[\*\-]\s+(.*)', self.lines[self.pos])
                    if lm:
                        indent = len(lm.group(1))
                        list_items.append({'indent': indent, 'text': lm.group(2).strip()})
                        self.pos += 1
                    elif self.lines[self.pos].strip() == '':
                        self.pos += 1
                        # Check if next non-empty line continues the list
                        if self.pos < len(self.lines) and re.match(r'^(\s*)[\*\-]\s+', self.lines[self.pos]):
                            continue
                        break
                    else:
                        break
                self.blocks.append({'type': 'list', 'items': list_items})
                continue
            
            # Ordered list
            if re.match(r'^\d+\.\s+', line):
                list_items = []
                while self.pos < len(self.lines):
                    lm = re.match(r'^(\d+)\.\s+(.*)', self.lines[self.pos])
                    if lm:
                        list_items.append({'num': int(lm.group(1)), 'text': lm.group(2).strip()})
                        self.pos += 1
                    elif self.lines[self.pos].strip() == '':
                        self.pos += 1
                        if self.pos < len(self.lines) and re.match(r'^\d+\.\s+', self.lines[self.pos]):
                            continue
                        break
                    else:
                        break
                self.blocks.append({'type': 'ordered_list', 'items': list_items})
                continue
            
            # Paragraph (collect consecutive lines)
            para_lines = []
            while self.pos < len(self.lines):
                pline = self.lines[self.pos]
                if (pline.strip() == '' or 
                    pline.startswith('#') or 
                    pline.startswith('```') or
                    pline.startswith('>') or
                    re.match(r'^---+\s*$', pline.strip()) or
                    re.match(r'^[\*\-]\s+', pline) or
                    re.match(r'^\d+\.\s+', pline) or
                    ('|' in pline and self.pos + 1 < len(self.lines) and 
                     re.match(r'^[\s|:-]+$', self.lines[self.pos + 1]))):
                    break
                para_lines.append(pline)
                self.pos += 1
            if para_lines:
                self.blocks.append({'type': 'paragraph', 'text': ' '.join(para_lines)})
    
    def _parse_table(self, lines):
        """Parse markdown table lines into structured data."""
        if len(lines) < 2:
            return
        
        def split_row(line):
            cells = [c.strip() for c in line.split('|')]
            # Remove empty first/last from leading/trailing |
            if cells and cells[0] == '':
                cells = cells[1:]
            if cells and cells[-1] == '':
                cells = cells[:-1]
            return cells
        
        headers = split_row(lines[0])
        # lines[1] is the separator
        rows = [split_row(l) for l in lines[2:] if l.strip()]
        
        self.blocks.append({'type': 'table', 'headers': headers, 'rows': rows})


# =========================================================================
# DOCX BUILDER
# =========================================================================

class DocxBuilder:
    """Build a DOCX document from parsed markdown blocks."""
    
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
        self._setup_page()
        self.temp_dir = tempfile.mkdtemp()
        self.mermaid_count = 0
    
    def _setup_page(self):
        """Configure page settings."""
        section = self.doc.sections[0]
        section.page_width = Cm(21.0)   # A4
        section.page_height = Cm(29.7)
        section.top_margin = PAGE_MARGIN
        section.bottom_margin = PAGE_MARGIN
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.54)
    
    def _setup_styles(self):
        """Set up document styles."""
        style = self.doc.styles['Normal']
        font = style.font
        font.name = FONT_MAIN
        font.size = Pt(FONT_SIZE_BODY)
        font.color.rgb = COLOR_TEXT
        
        pf = style.paragraph_format
        pf.space_after = Pt(6)
        pf.line_spacing = 1.15
    
    def _add_formatted_text(self, paragraph, text):
        """Add text with inline markdown formatting (bold, italic, code)."""
        # Process inline formatting: **bold**, *italic*, `code`, [links]
        pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|\[(.+?)\]\(.+?\))'
        
        pos = 0
        for match in re.finditer(pattern, text):
            # Add text before match
            before = text[pos:match.start()]
            if before:
                run = paragraph.add_run(before)
                run.font.name = FONT_MAIN
                run.font.size = Pt(FONT_SIZE_BODY)
            
            if match.group(2):  # ***bold italic***
                run = paragraph.add_run(match.group(2))
                run.bold = True
                run.italic = True
                run.font.name = FONT_MAIN
            elif match.group(3):  # **bold**
                run = paragraph.add_run(match.group(3))
                run.bold = True
                run.font.name = FONT_MAIN
            elif match.group(4):  # *italic*
                run = paragraph.add_run(match.group(4))
                run.italic = True
                run.font.name = FONT_MAIN
            elif match.group(5):  # `code`
                run = paragraph.add_run(match.group(5))
                run.font.name = FONT_MONO
                run.font.size = Pt(FONT_SIZE_BODY - 1)
                run.font.color.rgb = RGBColor(0xDB, 0x27, 0x77)
            elif match.group(6):  # [link text](url)
                run = paragraph.add_run(match.group(6))
                run.font.color.rgb = COLOR_PRIMARY
                run.font.name = FONT_MAIN
            
            pos = match.end()
        
        # Add remaining text
        remaining = text[pos:]
        if remaining:
            run = paragraph.add_run(remaining)
            run.font.name = FONT_MAIN
            run.font.size = Pt(FONT_SIZE_BODY)
    
    def _add_formatted_text_to_cell(self, paragraph, text, font_size=FONT_SIZE_TABLE):
        """Add formatted text in a table cell."""
        pattern = r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)'
        
        pos = 0
        for match in re.finditer(pattern, text):
            before = text[pos:match.start()]
            if before:
                run = paragraph.add_run(before)
                run.font.name = FONT_MAIN
                run.font.size = Pt(font_size)
            
            if match.group(2):
                run = paragraph.add_run(match.group(2))
                run.bold = True
                run.italic = True
                run.font.size = Pt(font_size)
            elif match.group(3):
                run = paragraph.add_run(match.group(3))
                run.bold = True
                run.font.size = Pt(font_size)
            elif match.group(4):
                run = paragraph.add_run(match.group(4))
                run.italic = True
                run.font.size = Pt(font_size)
            elif match.group(5):
                run = paragraph.add_run(match.group(5))
                run.font.name = FONT_MONO
                run.font.size = Pt(font_size - 1)
            
            run.font.name = FONT_MAIN
            pos = match.end()
        
        remaining = text[pos:]
        if remaining:
            run = paragraph.add_run(remaining)
            run.font.name = FONT_MAIN
            run.font.size = Pt(font_size)
    
    def add_heading(self, text, level):
        """Add a heading."""
        # Clean markdown formatting from heading text
        clean_text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)
        clean_text = re.sub(r'\*\*(.+?)\*\*', r'\1', clean_text)
        clean_text = re.sub(r'\*(.+?)\*', r'\1', clean_text)
        clean_text = re.sub(r'`(.+?)`', r'\1', clean_text)
        
        if level == 1:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_before = Pt(24)
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(clean_text.upper())
            run.bold = True
            run.font.size = Pt(FONT_SIZE_H1)
            run.font.color.rgb = COLOR_HEADING
            run.font.name = FONT_MAIN
            
            # Add underline bar
            p2 = self.doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(12)
            
        elif level == 2:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            run = p.add_run(clean_text)
            run.bold = True
            run.font.size = Pt(FONT_SIZE_H2)
            run.font.color.rgb = COLOR_PRIMARY
            run.font.name = FONT_MAIN
            
        elif level == 3:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(clean_text)
            run.bold = True
            run.font.size = Pt(FONT_SIZE_H3)
            run.font.color.rgb = COLOR_HEADING
            run.font.name = FONT_MAIN
            
        elif level >= 4:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(clean_text)
            run.bold = True
            run.font.size = Pt(FONT_SIZE_H4)
            run.font.color.rgb = COLOR_HEADING
            run.font.name = FONT_MAIN
    
    def add_paragraph(self, text):
        """Add a formatted paragraph."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        self._add_formatted_text(p, text)
    
    def add_blockquote(self, text):
        """Add a blockquote with left border styling."""
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        
        # Add left border via XML
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="12" w:space="8" w:color="D1D5DB"/></w:pBdr>')
        pPr.append(pBdr)
        
        self._add_formatted_text(p, text)
        for run in p.runs:
            run.italic = True
            run.font.color.rgb = COLOR_MUTED
    
    def add_alert(self, alert_type, text):
        """Add a GitHub-style alert box."""
        colors = ALERT_COLORS.get(alert_type, ALERT_COLORS['NOTE'])
        bg_color, border_color, icon = colors
        
        # Create a single-cell table to simulate a box
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        cell = table.cell(0, 0)
        
        # Set cell background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{bg_color}" w:val="clear"/>')
        tcPr.append(shading)
        
        # Set cell borders
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
            f'  <w:left w:val="single" w:sz="12" w:space="0" w:color="{border_color}"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
            f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{border_color}"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)
        
        # Alert type label
        p = cell.paragraphs[0]
        run = p.add_run(f"{icon}  {alert_type}")
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(border_color)
        run.font.name = FONT_MAIN
        
        # Alert text
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(4)
        self._add_formatted_text(p2, text)
        for run in p2.runs:
            run.font.size = Pt(FONT_SIZE_BODY - 1)
        
        # Add spacing after
        sp = self.doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(2)
    
    def add_table(self, headers, rows):
        """Add a formatted table."""
        num_cols = len(headers)
        table = self.doc.add_table(rows=1 + len(rows), cols=num_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True
        
        # Style headers
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self._add_formatted_text_to_cell(p, header, FONT_SIZE_TABLE)
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = COLOR_TABLE_HEADER_TEXT
            
            # Header background
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COLOR_TABLE_HEADER_BG}" w:val="clear"/>')
            tcPr.append(shading)
        
        # Style rows
        for r, row in enumerate(rows):
            for c in range(min(len(row), num_cols)):
                cell = table.cell(r + 1, c)
                p = cell.paragraphs[0]
                self._add_formatted_text_to_cell(p, row[c], FONT_SIZE_TABLE)
                
                # Alternate row shading
                if r % 2 == 1:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COLOR_TABLE_ALT_ROW}" w:val="clear"/>')
                    tcPr.append(shading)
        
        # Set border style for whole table
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
            f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
            f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
            f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
            f'</w:tblBorders>'
        )
        tblPr.append(borders)
        
        # Spacing after table
        sp = self.doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(4)
    
    def add_list(self, items):
        """Add an unordered list."""
        for item in items:
            indent_level = item.get('indent', 0) // 2
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.5 + indent_level * 0.7)
            
            # Bullet symbol
            bullet = '•' if indent_level == 0 else '◦'
            run = p.add_run(f"{bullet}  ")
            run.font.name = FONT_MAIN
            run.font.size = Pt(FONT_SIZE_BODY)
            
            self._add_formatted_text(p, item['text'])
    
    def add_ordered_list(self, items):
        """Add an ordered list."""
        for item in items:
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Cm(0.5)
            
            run = p.add_run(f"{item['num']}.  ")
            run.font.name = FONT_MAIN
            run.font.size = Pt(FONT_SIZE_BODY)
            run.bold = True
            
            self._add_formatted_text(p, item['text'])
    
    def add_code_block(self, code, lang=''):
        """Add a code block with background shading."""
        # Create a single-cell table for the code box
        table = self.doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        cell = table.cell(0, 0)
        
        # Set cell background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COLOR_CODE_BG}" w:val="clear"/>')
        tcPr.append(shading)
        
        # Set border
        borders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
            f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
            f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="D1D5DB"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(borders)
        
        p = cell.paragraphs[0]
        for i, line in enumerate(code.split('\n')):
            if i > 0:
                p.add_run('\n')
            run = p.add_run(line)
            run.font.name = FONT_MONO
            run.font.size = Pt(FONT_SIZE_SMALL)
            run.font.color.rgb = COLOR_TEXT
        
        sp = self.doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(4)
    
    def add_mermaid(self, code):
        """Render and add a Mermaid diagram as an image."""
        self.mermaid_count += 1
        img_path = os.path.join(self.temp_dir, f'mermaid_{self.mermaid_count}.png')
        
        print(f"  📊 Rendering Mermaid diagram #{self.mermaid_count}...")
        
        if render_mermaid_to_png(code, img_path):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(12)
            
            # Calculate appropriate width (max 15cm for A4)
            run = p.add_run()
            try:
                run.add_picture(img_path, width=Cm(15))
            except Exception:
                # If the image is too wide, try smaller
                try:
                    run.add_picture(img_path, width=Cm(12))
                except Exception as e:
                    run.text = f"[Diagram Mermaid #{self.mermaid_count} — gagal dimuat: {e}]"
                    run.font.color.rgb = COLOR_MUTED
                    run.italic = True
            
            print(f"  ✅ Diagram #{self.mermaid_count} berhasil di-embed")
        else:
            # Fallback: show as code block
            print(f"  ⚠️  Fallback: menampilkan diagram #{self.mermaid_count} sebagai kode")
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(f"[Diagram Mermaid #{self.mermaid_count}]")
            run.font.color.rgb = COLOR_MUTED
            run.italic = True
            run.font.size = Pt(FONT_SIZE_SMALL)
            self.add_code_block(code, 'mermaid')
    
    def add_hr(self):
        """Add a horizontal rule."""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)
        
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(
            f'<w:pBdr {nsdecls("w")}>'
            f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="D1D5DB"/>'
            f'</w:pBdr>'
        )
        pPr.append(pBdr)
    
    def add_cover_page(self, title, subtitle, organization, date_str):
        """Add a professional cover page."""
        # Add several blank lines for spacing
        for _ in range(6):
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
        
        # Title
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(28)
        run.font.color.rgb = COLOR_HEADING
        run.font.name = FONT_MAIN
        
        # Subtitle
        if subtitle:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            run = p.add_run(subtitle)
            run.italic = True
            run.font.size = Pt(16)
            run.font.color.rgb = COLOR_PRIMARY
            run.font.name = FONT_MAIN
        
        # Divider
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(20)
        run = p.add_run('━' * 30)
        run.font.color.rgb = COLOR_PRIMARY
        run.font.size = Pt(12)
        
        # Extra info lines
        for line in organization:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line)
            run.font.size = Pt(12)
            run.font.color.rgb = COLOR_TEXT
            run.font.name = FONT_MAIN
        
        # Date
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(20)
        run = p.add_run(date_str)
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = COLOR_MUTED
        run.font.name = FONT_MAIN
        
        # Page break
        self.doc.add_page_break()
    
    def build(self, blocks):
        """Build the document from parsed blocks."""
        # Check if first block is the title
        skip_cover = False
        if (blocks and blocks[0]['type'] == 'heading' and blocks[0]['level'] == 1):
            title = blocks[0]['text']
            # Extract subtitle and org info from the next few blocks
            subtitle = ''
            org_lines = []
            date_str = ''
            
            i = 1
            while i < len(blocks) and blocks[i]['type'] in ('heading', 'paragraph', 'hr'):
                b = blocks[i]
                if b['type'] == 'heading' and b['level'] == 3:
                    subtitle = re.sub(r'\*(.+?)\*', r'\1', b['text'])
                elif b['type'] == 'paragraph':
                    text = b['text'].strip()
                    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
                    if 'Diajukan' in text or 'Disusun' in text or 'Balai' in text or 'Kementerian' in text:
                        org_lines.append(text)
                    elif re.match(r'.*\d{4}.*', text) and len(text) < 30:
                        date_str = text
                    else:
                        org_lines.append(text)
                elif b['type'] == 'hr':
                    i += 1
                    break
                i += 1
            
            if not date_str:
                date_str = 'Juni 2026'
            
            self.add_cover_page(title, subtitle, org_lines, date_str)
            skip_cover = True
            blocks = blocks[i:]
        
        for block in blocks:
            btype = block['type']
            
            if btype == 'heading':
                self.add_heading(block['text'], block['level'])
            elif btype == 'paragraph':
                self.add_paragraph(block['text'])
            elif btype == 'blockquote':
                self.add_blockquote(block['text'])
            elif btype == 'alert':
                self.add_alert(block['alert_type'], block['text'])
            elif btype == 'table':
                self.add_table(block['headers'], block['rows'])
            elif btype == 'list':
                self.add_list(block['items'])
            elif btype == 'ordered_list':
                self.add_ordered_list(block['items'])
            elif btype == 'mermaid':
                self.add_mermaid(block['code'])
            elif btype == 'code':
                self.add_code_block(block['code'], block.get('lang', ''))
            elif btype == 'hr':
                self.add_hr()
    
    def save(self, output_path):
        """Save the document."""
        self.doc.save(output_path)
        
        # Cleanup temp files
        import shutil
        try:
            shutil.rmtree(self.temp_dir, ignore_errors=True)
        except:
            pass


# =========================================================================
# MAIN
# =========================================================================

def convert_md_to_docx(input_path: str, output_path: str):
    """Convert a markdown file to DOCX."""
    print(f"\n{'='*60}")
    print(f"  MD → DOCX Converter with Mermaid Support")
    print(f"{'='*60}")
    print(f"\n  📄 Input:  {input_path}")
    print(f"  📦 Output: {output_path}\n")
    
    # Read markdown
    with open(input_path, 'r', encoding='utf-8') as f:
        md_text = f.read()
    
    print(f"  📝 Parsing markdown ({len(md_text)} characters)...")
    
    # Parse
    parser = MarkdownParser(md_text)
    print(f"  ✅ Parsed {len(parser.blocks)} blocks")
    
    # Count mermaid diagrams
    mermaid_count = sum(1 for b in parser.blocks if b['type'] == 'mermaid')
    if mermaid_count:
        print(f"  🎨 Found {mermaid_count} Mermaid diagram(s) to render")
    
    # Build DOCX
    print(f"\n  🔨 Building DOCX document...")
    builder = DocxBuilder()
    builder.build(parser.blocks)
    
    # Save
    builder.save(output_path)
    
    file_size = os.path.getsize(output_path)
    print(f"\n  ✅ DOCX berhasil dibuat!")
    print(f"  📦 File: {output_path}")
    print(f"  📏 Size: {file_size / 1024:.1f} KB")
    print(f"{'='*60}\n")


if __name__ == '__main__':
    # Default paths
    default_input = os.path.join(
        os.path.expanduser('~'), '.gemini', 'antigravity-ide', 'brain',
        '6f99d702-2f67-42c4-8812-85aca1761be5', 'design_thinking_adaptika.md'
    )
    
    # Also check Downloads
    downloads_input = os.path.join(
        os.path.expanduser('~'), 'Downloads', 'design_thinking_adaptika.md'
    )
    
    script_dir = os.path.dirname(os.path.abspath(__file__))
    default_output = os.path.join(script_dir, 'Makalah_Inovasi_ADAPTIKA.docx')
    
    if len(sys.argv) >= 3:
        input_path = sys.argv[1]
        output_path = sys.argv[2]
    elif len(sys.argv) == 2:
        input_path = sys.argv[1]
        output_path = default_output
    else:
        # Try default locations
        if os.path.exists(default_input):
            input_path = default_input
        elif os.path.exists(downloads_input):
            input_path = downloads_input
        else:
            print(f"❌ File not found at default locations:")
            print(f"   {default_input}")
            print(f"   {downloads_input}")
            print(f"\nUsage: python md_to_docx.py [input.md] [output.docx]")
            sys.exit(1)
        output_path = default_output
    
    if not os.path.exists(input_path):
        print(f"❌ Input file not found: {input_path}")
        sys.exit(1)
    
    convert_md_to_docx(input_path, output_path)
